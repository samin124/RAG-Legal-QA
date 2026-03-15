"""
PDF Loading and Chunking
Handles document parsing with LlamaParse (OCR support) and text splitting
"""

from typing import List, Optional
import os
import tempfile
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document


def load_pdf_with_llamaparse(
    file_path: str,
    use_ocr: bool = True
) -> List[Document]:
    """
    Load and parse a PDF document using LlamaParse with OCR support.

    Args:
        file_path: Path to the PDF file
        use_ocr: Enable OCR for scanned documents

    Returns:
        List of Document objects
    """
    from llama_parse import LlamaParse

    # Configure LlamaParse with OCR and page-aware splitting
    parser = LlamaParse(
        api_key=os.getenv("LLAMA_CLOUD_API_KEY"),
        result_type="markdown",
        parsing_instruction="Extract all text from this legal document. Pay attention to clauses, sections, and legal terminology. Preserve page boundaries.",
        use_vendor_multimodal_model=use_ocr,
        vendor_multimodal_model_name="gemini-2.0-flash" if use_ocr else None,
        split_by_page=True,  # This ensures each page becomes a separate document
        page_separator="\n---PAGE_BREAK---\n"
    )

    # Parse the document
    documents = parser.load_data(file_path)

    print(f"DEBUG: LlamaParse returned {len(documents)} document(s)")

    # Convert to LangChain Document format
    langchain_docs = []

    # If LlamaParse returned a single document, try to split by page separator
    if len(documents) == 1 and "---PAGE_BREAK---" in documents[0].text:
        print("DEBUG: Detected page breaks, splitting single document into pages")
        pages = documents[0].text.split("---PAGE_BREAK---")
        for page_num, page_content in enumerate(pages, start=1):
            if page_content.strip():  # Skip empty pages
                print(f"DEBUG: Split page {page_num}, Content length: {len(page_content)}")
                langchain_docs.append(Document(
                    page_content=page_content.strip(),
                    metadata={
                        "source": file_path,
                        "page": page_num,
                        "parser": "llamaparse",
                        "ocr_enabled": use_ocr,
                        "split_method": "page_separator"
                    }
                ))
    else:
        # Standard processing - one document per page
        for i, doc in enumerate(documents):
            # Try to get page number from LlamaParse metadata if available
            page_num = i + 1
            if hasattr(doc, 'metadata') and 'page' in doc.metadata:
                page_num = doc.metadata['page']
            elif hasattr(doc, 'extra_info') and 'page' in doc.extra_info:
                page_num = doc.extra_info['page']

            print(f"DEBUG: Document {i} -> Page {page_num}, Content length: {len(doc.text)}")

            langchain_docs.append(Document(
                page_content=doc.text,
                metadata={
                    "source": file_path,
                    "page": page_num,
                    "parser": "llamaparse",
                    "ocr_enabled": use_ocr,
                    "doc_index": i
                }
            ))

    print(f"DEBUG: Final document count: {len(langchain_docs)} pages")
    return langchain_docs


def load_pdf_from_bytes(
    file_bytes: bytes,
    filename: str,
    use_ocr: bool = True
) -> List[Document]:
    """
    Load PDF from uploaded bytes (for Streamlit uploads).

    Args:
        file_bytes: Raw PDF file bytes
        filename: Original filename for metadata
        use_ocr: Enable OCR for scanned documents

    Returns:
        List of Document objects
    """
    # Write to temp file for LlamaParse
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    try:
        documents = load_pdf_with_llamaparse(tmp_path, use_ocr=use_ocr)
        # Update source metadata with original filename
        for doc in documents:
            doc.metadata["source"] = filename
            doc.metadata["original_path"] = tmp_path
        return documents
    finally:
        # Clean up temp file
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


def load_pdf_fallback(file_path: str) -> List[Document]:
    """
    Fallback PDF loading using PyPDF (when LlamaParse is unavailable).

    Args:
        file_path: Path to the PDF file

    Returns:
        List of Document objects, one per page
    """
    from langchain_community.document_loaders import PyPDFLoader

    loader = PyPDFLoader(file_path)
    documents = loader.load()

    print(f"DEBUG: PyPDF loaded {len(documents)} pages")

    # Add parser info to metadata and ensure page numbers are correct
    for i, doc in enumerate(documents):
        doc.metadata["parser"] = "pypdf"
        doc.metadata["ocr_enabled"] = False
        # PyPDF usually provides page numbers, but verify
        if "page" not in doc.metadata:
            doc.metadata["page"] = i + 1
        print(f"DEBUG: PyPDF page {i+1}, metadata page: {doc.metadata.get('page')}")

    return documents


def chunk_documents(
    documents: List[Document],
    chunk_size: int = 1000,
    chunk_overlap: int = 200
) -> List[Document]:
    """
    Split documents into overlapping chunks using RecursiveCharacterTextSplitter.

    Args:
        documents: List of Document objects to split
        chunk_size: Target size for each chunk (in characters)
        chunk_overlap: Overlap between consecutive chunks

    Returns:
        List of chunked Document objects
    """
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        separators=[
            "\n\n",      # Paragraph breaks
            "\n",        # Line breaks
            ". ",        # Sentence endings
            "; ",        # Clause separators
            ", ",        # Comma breaks
            " ",         # Word breaks
            ""           # Character level
        ]
    )

    chunks = text_splitter.split_documents(documents)

    print(f"DEBUG: Created {len(chunks)} chunks from {len(documents)} documents")

    # Add chunk index to metadata and verify page numbers
    page_distribution = {}
    for i, chunk in enumerate(chunks):
        chunk.metadata["chunk_index"] = i
        chunk.metadata["total_chunks"] = len(chunks)

        # Track page distribution for debugging
        page_num = chunk.metadata.get("page", "UNKNOWN")
        page_distribution[page_num] = page_distribution.get(page_num, 0) + 1

    print(f"DEBUG: Chunk page distribution: {page_distribution}")

    return chunks


def process_uploaded_pdf(
    file_bytes: bytes,
    filename: str,
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
    use_ocr: bool = True
) -> List[Document]:
    """
    Complete pipeline: Load PDF and chunk it.

    Args:
        file_bytes: Raw PDF file bytes
        filename: Original filename
        chunk_size: Target chunk size
        chunk_overlap: Chunk overlap
        use_ocr: Enable OCR for scanned documents

    Returns:
        List of chunked Document objects ready for embedding
    """
    try:
        # Try LlamaParse first
        documents = load_pdf_from_bytes(file_bytes, filename, use_ocr=use_ocr)
    except Exception as e:
        # Fallback to PyPDF if LlamaParse fails
        print(f"LlamaParse failed, falling back to PyPDF: {e}")
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
        try:
            documents = load_pdf_fallback(tmp_path)
            for doc in documents:
                doc.metadata["source"] = filename
        finally:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)

    # Chunk the documents
    chunks = chunk_documents(
        documents,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap
    )

    return chunks


def load_image_from_bytes(
    file_bytes: bytes,
    filename: str,
    use_ocr: bool = True
) -> List[Document]:
    """
    Load image file and extract text using OCR (via LlamaParse).

    Args:
        file_bytes: Raw image file bytes
        filename: Original filename
        use_ocr: Enable OCR for text extraction

    Returns:
        List of Document objects
    """
    # Write to temp file
    suffix = os.path.splitext(filename)[1] or ".png"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    try:
        # Use LlamaParse for image OCR
        from llama_parse import LlamaParse

        parser = LlamaParse(
            api_key=os.getenv("LLAMA_CLOUD_API_KEY"),
            result_type="markdown",
            parsing_instruction="Extract all text from this image document. Pay attention to any legal content, clauses, or formal text.",
            use_vendor_multimodal_model=True,
            vendor_multimodal_model_name="gemini-2.0-flash",
        )

        documents = parser.load_data(tmp_path)

        print(f"DEBUG: LlamaParse (image) returned {len(documents)} document(s)")

        # Convert to LangChain format
        langchain_docs = []
        for i, doc in enumerate(documents):
            # For images, page is usually 1 unless multi-page
            page_num = i + 1
            if hasattr(doc, 'metadata') and 'page' in doc.metadata:
                page_num = doc.metadata['page']

            print(f"DEBUG: Image document {i} -> Page {page_num}")

            langchain_docs.append(Document(
                page_content=doc.text,
                metadata={
                    "source": filename,
                    "page": page_num,
                    "parser": "llamaparse_image",
                    "ocr_enabled": use_ocr,
                    "file_type": "image"
                }
            ))

        return langchain_docs
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


def load_docx_from_bytes(
    file_bytes: bytes,
    filename: str
) -> List[Document]:
    """
    Load Word document (.docx) and extract text.

    Args:
        file_bytes: Raw DOCX file bytes
        filename: Original filename

    Returns:
        List of Document objects
    """
    from docx import Document as DocxDocument
    import io

    # Load from bytes
    doc = DocxDocument(io.BytesIO(file_bytes))

    # Extract all paragraphs
    text_content = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_content.append(para.text)

    # Combine into single document
    full_text = "\n\n".join(text_content)

    return [Document(
        page_content=full_text,
        metadata={
            "source": filename,
            "parser": "python-docx",
            "file_type": "docx",
            "paragraph_count": len(text_content)
        }
    )]


def load_markdown_from_bytes(
    file_bytes: bytes,
    filename: str
) -> List[Document]:
    """
    Load Markdown (.md) file.

    Args:
        file_bytes: Raw MD file bytes
        filename: Original filename

    Returns:
        List of Document objects
    """
    # Decode text
    text_content = file_bytes.decode('utf-8')

    return [Document(
        page_content=text_content,
        metadata={
            "source": filename,
            "parser": "markdown",
            "file_type": "md"
        }
    )]


def process_uploaded_document(
    file_bytes: bytes,
    filename: str,
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
    use_ocr: bool = True
) -> List[Document]:
    """
    Universal document processor: handles PDF, images, DOCX, and MD files.

    Args:
        file_bytes: Raw file bytes
        filename: Original filename
        chunk_size: Target chunk size
        chunk_overlap: Chunk overlap
        use_ocr: Enable OCR for images and scanned PDFs

    Returns:
        List of chunked Document objects ready for embedding
    """
    # Determine file type from extension
    file_ext = os.path.splitext(filename)[1].lower()

    # Route to appropriate loader
    if file_ext == '.pdf':
        return process_uploaded_pdf(file_bytes, filename, chunk_size, chunk_overlap, use_ocr)

    elif file_ext in ['.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.gif']:
        documents = load_image_from_bytes(file_bytes, filename, use_ocr)

    elif file_ext == '.docx':
        documents = load_docx_from_bytes(file_bytes, filename)

    elif file_ext in ['.md', '.markdown']:
        documents = load_markdown_from_bytes(file_bytes, filename)

    else:
        raise ValueError(f"Unsupported file type: {file_ext}. Supported types: PDF, images (PNG, JPG), DOCX, MD")

    # Chunk the documents
    chunks = chunk_documents(
        documents,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap
    )

    return chunks
